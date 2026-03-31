"""
Report Generator - 정보보호공시 경영진 보고용 리포트
KISA 정보보호 공시 가이드라인(2024.03) 기반
python-docx를 사용하여 전문적인 Word 문서 생성
"""
import io
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ═══════════════════════════════════
# 색상 팔레트
# ═══════════════════════════════════
C_NAVY      = "1E3A5F"
C_DARK_NAVY = "0F1F33"
C_LIGHT_NAVY= "2C5282"
C_ACCENT    = "3182CE"   # 파란 강조
C_GREEN     = "276749"
C_L_GREEN   = "E6F4EA"
C_PURPLE    = "5E35B1"
C_L_PURPLE  = "EDE7F6"
C_ORANGE    = "DD6B20"
C_L_ORANGE  = "FFF5EB"
C_RED       = "C53030"
C_L_RED     = "FEE2E2"
C_GRAY      = "666666"
C_L_GRAY    = "F7FAFC"
C_HEADER_BG = "1E3A5F"
C_ROW_ALT   = "EDF2F7"
C_WHITE     = "FFFFFF"
C_BLACK     = "1A202C"


def _rgb(hex_str):
    return RGBColor(int(hex_str[:2], 16), int(hex_str[2:4], 16), int(hex_str[4:], 16))


# ═══════════════════════════════════
# 셀/단락 스타일 헬퍼
# ═══════════════════════════════════
def _set_cell_shading(cell, color_hex):
    tc_pr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color_hex}"/>')
    tc_pr.append(shd)


def _set_cell_border(cell, **kwargs):
    """셀 테두리 설정: top, bottom, left, right"""
    tc_pr = cell._element.get_or_add_tcPr()
    borders = tc_pr.find(qn('w:tcBorders'))
    if borders is None:
        borders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
        tc_pr.append(borders)
    for edge, val in kwargs.items():
        edge_el = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val","single")}" '
            f'w:sz="{val.get("sz","4")}" w:space="0" w:color="{val.get("color","auto")}"/>'
        )
        existing = borders.find(qn(f'w:{edge}'))
        if existing is not None:
            borders.remove(existing)
        borders.append(edge_el)


def _add_paragraph_border_bottom(paragraph, color=C_NAVY, size="12"):
    """단락 하단 테두리선"""
    pPr = paragraph._element.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(borders)


def _make_run(paragraph, text, size=10, bold=False, color=C_BLACK, italic=False, font="Malgun Gothic"):
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = _rgb(color)
    run.font.italic = italic
    run.font.name = font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    return run


# ═══════════════════════════════════
# 전문 테이블 생성기
# ═══════════════════════════════════
def _add_exec_table(doc, headers, rows, col_widths=None, header_color=C_HEADER_BG,
                    accent_col=None, total_row=False):
    """경영진 보고용 세련된 테이블"""
    n_rows = 1 + len(rows)
    n_cols = len(headers)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 테이블 전체 테두리 제거 후 커스텀 적용
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')

    # 헤더 행
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _make_run(p, str(h), size=9, bold=True, color=C_WHITE)
        _set_cell_shading(cell, header_color)
        # 헤더 상하단 두꺼운 테두리
        _set_cell_border(cell,
            top={"val": "single", "sz": "8", "color": header_color},
            bottom={"val": "single", "sz": "8", "color": header_color},
        )

    # 데이터 행
    for r_idx, row_data in enumerate(rows):
        is_last = r_idx == len(rows) - 1
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            val_str = str(val)

            # 숫자 우측 정렬
            if c_idx > 0 and _is_numeric_str(val_str):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            is_bold = (total_row and is_last) or (accent_col is not None and c_idx == accent_col)
            text_color = C_NAVY if is_bold else C_BLACK
            _make_run(p, val_str, size=9, bold=is_bold, color=text_color)

            # 교차 행 배경
            if r_idx % 2 == 1 and not (total_row and is_last):
                _set_cell_shading(cell, C_ROW_ALT)

            # 총계 행 강조
            if total_row and is_last:
                _set_cell_shading(cell, "E2E8F0")
                _set_cell_border(cell, top={"val": "single", "sz": "6", "color": C_NAVY})

            # 하단 얇은 테두리
            _set_cell_border(cell, bottom={"val": "single", "sz": "2", "color": "E2E8F0"})

    # 열 너비
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    doc.add_paragraph("")  # 테이블 후 여백
    return table


def _is_numeric_str(s):
    try:
        s = s.replace(",", "").replace("%", "").replace("원", "").replace("억", "").replace("만", "").strip()
        if not s or s == "-":
            return False
        float(s)
        return True
    except:
        return False


# ═══════════════════════════════════
# 강조 박스 (Executive Callout)
# ═══════════════════════════════════
def _add_callout_box(doc, items, bg_color=C_L_GRAY, border_color=C_NAVY):
    """핵심 수치 강조 박스"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, bg_color)
    for edge in ["top", "bottom", "left", "right"]:
        _set_cell_border(cell, **{edge: {"val": "single", "sz": "4", "color": border_color}})
    # 왼쪽 두꺼운 강조선
    _set_cell_border(cell, left={"val": "single", "sz": "18", "color": border_color})

    for i, (label, value) in enumerate(items):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.space_after = Pt(4)
        _make_run(p, f"  {label}: ", size=10, bold=False, color=C_GRAY)
        _make_run(p, str(value), size=12, bold=True, color=C_NAVY)

    doc.add_paragraph("")


# ═══════════════════════════════════
# 섹션 헤딩 (번호 + 밑줄)
# ═══════════════════════════════════
def _add_section_heading(doc, number, title):
    """섹션 제목 - 남색 번호 + 하단 구분선"""
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    p.space_after = Pt(8)
    _make_run(p, f"{number}  ", size=16, bold=True, color=C_ACCENT)
    _make_run(p, title, size=16, bold=True, color=C_DARK_NAVY)
    _add_paragraph_border_bottom(p, color=C_ACCENT, size="8")
    return p


def _add_sub_heading(doc, title, color=C_LIGHT_NAVY):
    """서브 제목"""
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(4)
    _make_run(p, "▎ ", size=11, bold=True, color=color)
    _make_run(p, title, size=11, bold=True, color=C_DARK_NAVY)
    return p


def _add_body(doc, text, size=10, color=C_BLACK, bold=False, indent=False, italic=False):
    """본문 텍스트"""
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    _make_run(p, text, size=size, bold=bold, color=color, italic=italic)
    return p


def _fmt_amount(val):
    if pd.isna(val) or val == "":
        return "0"
    try:
        v = float(val)
        if abs(v) >= 1e8:
            return f"{v/1e8:,.1f}억"
        elif abs(v) >= 1e4:
            return f"{v/1e4:,.0f}만"
        else:
            return f"{v:,.0f}"
    except:
        return str(val)


def _fmt_won(val):
    """원 단위 포맷"""
    try:
        return f"{abs(float(val)):,.0f}원"
    except:
        return "0원"


# ═══════════════════════════════════
# 메인 리포트 생성 함수
# ═══════════════════════════════════
def generate_report(asset_df: pd.DataFrame, cost_df: pd.DataFrame) -> io.BytesIO:
    doc = Document()
    today = datetime.now()

    # ── 페이지/기본 스타일 ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    style.paragraph_format.space_after = Pt(6)

    # ── 헤더 (2페이지부터) ──
    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _make_run(hp, "정보보호공시 보고서  |  ", size=7, color=C_GRAY, italic=True)
    _make_run(hp, "CONFIDENTIAL", size=7, color=C_RED, bold=True, italic=True)

    # ── 푸터 (페이지 번호) ──
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _make_run(fp, "AI기반 정보보호공시 솔루션  |  ", size=7, color=C_GRAY)
    # 페이지 번호 필드
    run = fp.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)
    run2 = fp.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._element.append(instrText)
    run3 = fp.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fldChar2)

    # ── 데이터 집계 ──
    stats = _compute_stats(asset_df, cost_df)
    it_total = stats["asset_it_total"] + stats["cost_it_total"]
    sec_total = stats["asset_sec_total"] + stats["cost_sec_total"]
    ratio = (sec_total / it_total * 100) if it_total > 0 else 0

    # ═══════════════════════════════════
    # 표지
    # ═══════════════════════════════════
    for _ in range(3):
        doc.add_paragraph("")

    # 상단 장식선
    line1 = doc.add_paragraph()
    line1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _make_run(line1, "━" * 45, size=10, color=C_ACCENT)

    for _ in range(2):
        doc.add_paragraph("")

    cover_t = doc.add_paragraph()
    cover_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_t.space_after = Pt(4)
    _make_run(cover_t, "정보보호공시 보고서", size=38, bold=True, color=C_DARK_NAVY)

    doc.add_paragraph("")

    cover_sub = doc.add_paragraph()
    cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _make_run(cover_sub, "Information Security Disclosure Report", size=14, color=C_GRAY, italic=True)

    doc.add_paragraph("")

    cover_desc = doc.add_paragraph()
    cover_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _make_run(cover_desc, "KISA 정보보호 공시 가이드라인(2024.03) 기반", size=12, color=C_LIGHT_NAVY)

    cover_desc2 = doc.add_paragraph()
    cover_desc2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _make_run(cover_desc2, "정보기술부문 / 정보보호부문 투자 분석", size=12, color=C_LIGHT_NAVY)

    for _ in range(4):
        doc.add_paragraph("")

    # 하단 장식선
    line2 = doc.add_paragraph()
    line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _make_run(line2, "━" * 45, size=10, color=C_ACCENT)

    doc.add_paragraph("")

    # 날짜 & 기밀
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _make_run(date_p, today.strftime("%Y년 %m월 %d일"), size=14, bold=True, color=C_NAVY)

    doc.add_paragraph("")

    conf_p = doc.add_paragraph()
    conf_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _make_run(conf_p, "CONFIDENTIAL", size=11, bold=True, color=C_RED)

    gen_p = doc.add_paragraph()
    gen_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _make_run(gen_p, "AI기반 정보보호공시 솔루션  |  자동 생성 문서", size=9, color=C_GRAY)

    doc.add_page_break()

    # ═══════════════════════════════════
    # 목차
    # ═══════════════════════════════════
    toc_h = doc.add_paragraph()
    toc_h.space_after = Pt(16)
    _make_run(toc_h, "목  차", size=24, bold=True, color=C_DARK_NAVY)
    _add_paragraph_border_bottom(toc_h, color=C_ACCENT, size="12")

    doc.add_paragraph("")

    toc_items = [
        ("01", "Executive Summary", "총 IT 투자 / 정보보호 투자 요약 및 핵심 지표"),
        ("02", "투자 현황 Overview", "자산·비용별 정보기술 / 정보보호 / 제외 투자 현황"),
        ("03", "세부 분류별 투자 Breakdown", "자산분류별 · 계정별 상세 내역"),
        ("04", "주요 투자 항목 분석 (Top 10)", "금액 기준 상위 자산 · 비용 항목"),
        ("05", "분류 기준 및 판단 로직", "KISA 가이드라인 기준 · 자동 분류 로직 (Audit 대응)"),
        ("06", "회계 정합성 및 산출 방법", "데이터 출처 · 산출 기준 · 중복 제거 로직"),
        ("07", "분석 매칭 결과 통계", "완전일치 / 포함매칭 / 유사매칭 / AI분류 / 미매칭"),
        ("08", "리스크 및 개선사항", "보안 투자 비율 분석 · 미분류 항목 검토"),
    ]

    for num, title, desc in toc_items:
        p = doc.add_paragraph()
        p.space_after = Pt(2)
        _make_run(p, f"  {num}   ", size=12, bold=True, color=C_ACCENT)
        _make_run(p, title, size=11, bold=True, color=C_DARK_NAVY)
        # 설명
        p2 = doc.add_paragraph()
        p2.space_after = Pt(10)
        p2.paragraph_format.left_indent = Cm(1.5)
        _make_run(p2, desc, size=9, color=C_GRAY)

    doc.add_paragraph("")
    ref_p = doc.add_paragraph()
    _make_run(ref_p, "※ 본 보고서는 KISA 「정보보호 공시 가이드라인(2024.03 개정본)」을 기준으로 작성되었습니다.",
              size=8, color=C_GRAY, italic=True)

    doc.add_page_break()

    # ═══════════════════════════════════
    # ① Executive Summary
    # ═══════════════════════════════════
    _add_section_heading(doc, "01", "Executive Summary")

    _add_body(doc,
        "본 리포트는 KISA 「정보보호 공시 가이드라인(2024.03 개정본)」에 따라 "
        "자산대장 및 비용원장의 분류 결과를 분석한 정보보호 공시용 보고서입니다. "
        "정보기술부문과 정보보호부문의 투자 현황을 정량적으로 분석하여 제시합니다."
    )

    doc.add_paragraph("")

    # 핵심 지표 강조 박스
    _add_callout_box(doc, [
        ("정보기술(IT) 총 투자", _fmt_won(it_total)),
        ("정보보호 총 투자", _fmt_won(sec_total)),
        ("정보보호 / IT 투자 비율", f"{ratio:.1f}%"),
    ], bg_color="EBF4FF", border_color=C_ACCENT)

    # 상세 요약 테이블
    _add_sub_heading(doc, "부문별 투자 현황 요약")
    _add_exec_table(doc,
        ["구분", "금액", "건수"],
        [
            ["정보기술 투자 (자산)", _fmt_won(stats["asset_it_total"]), f'{stats["asset_it_cnt"]}건'],
            ["정보기술 투자 (비용)", _fmt_won(stats["cost_it_total"]), f'{stats["cost_it_cnt"]}건'],
            ["정보보호 투자 (자산)", _fmt_won(stats["asset_sec_total"]), f'{stats["asset_sec_cnt"]}건'],
            ["정보보호 투자 (비용)", _fmt_won(stats["cost_sec_total"]), f'{stats["cost_sec_cnt"]}건'],
            ["제외 항목 (자산)", _fmt_won(stats["asset_exc_total"]), f'{stats["asset_exc_cnt"]}건'],
            ["제외 항목 (비용)", _fmt_won(stats["cost_exc_total"]), f'{stats["cost_exc_cnt"]}건'],
        ],
        col_widths=[6, 5, 4]
    )

    doc.add_page_break()

    # ═══════════════════════════════════
    # ② 투자 현황 Overview
    # ═══════════════════════════════════
    _add_section_heading(doc, "02", "투자 현황 Overview")

    _add_exec_table(doc,
        ["구분", "자산(감가상각비)", "비용(발생액)", "합계", "비중"],
        [
            ["정보기술(IT) 총 투자",
             _fmt_won(stats["asset_it_total"]),
             _fmt_won(stats["cost_it_total"]),
             _fmt_won(it_total), "100%"],
            ["  └ 정보보호 투자",
             _fmt_won(stats["asset_sec_total"]),
             _fmt_won(stats["cost_sec_total"]),
             _fmt_won(sec_total), f"{ratio:.1f}%"],
            ["제외",
             _fmt_won(stats["asset_exc_total"]),
             _fmt_won(stats["cost_exc_total"]),
             _fmt_won(stats["asset_exc_total"] + stats["cost_exc_total"]), "-"],
        ],
        col_widths=[4, 3, 3, 3, 2],
        total_row=False
    )

    _add_body(doc, "※ 정보보호 투자는 정보기술(IT) 투자의 부분집합입니다. (IT ⊃ 정보보호)", size=9, color=C_GRAY)
    _add_body(doc, "※ 자산 투자액은 연중 감가상각비 기준, 비용은 발생주의 기준으로 산출합니다.", size=9, color=C_GRAY)

    # 비율 강조
    if ratio > 0:
        level = "양호" if ratio >= 10 else ("보통" if ratio >= 5 else "미흡")
        level_color = C_L_GREEN if ratio >= 10 else (C_L_ORANGE if ratio >= 5 else C_L_RED)
        border_c = C_GREEN if ratio >= 10 else (C_ORANGE if ratio >= 5 else C_RED)
        _add_callout_box(doc, [
            ("정보보호 투자 비율 평가", f"{ratio:.1f}% — {level} (업계 권장 5~15%)"),
        ], bg_color=level_color, border_color=border_c)

    doc.add_page_break()

    # ═══════════════════════════════════
    # ③ 세부 분류별 Breakdown
    # ═══════════════════════════════════
    _add_section_heading(doc, "03", "세부 분류별 투자 Breakdown")

    _add_sub_heading(doc, "자산 분류별 현황", color=C_GREEN)
    if not asset_df.empty and "asset_class" in asset_df.columns:
        rows = _group_by_class(asset_df, "asset_class", "depreciation")
        if rows:
            _add_exec_table(doc,
                ["자산분류", "건수", "감가상각비", "IT", "보안"],
                rows, col_widths=[5, 2, 3, 2, 2], header_color=C_GREEN
            )
        else:
            _add_body(doc, "자산 분류별 데이터가 없습니다.", color=C_GRAY)
    else:
        _add_body(doc, "자산 분석 데이터가 없습니다.", color=C_GRAY)

    _add_sub_heading(doc, "비용 계정별 현황", color=C_PURPLE)
    if not cost_df.empty and "account_name" in cost_df.columns:
        rows = _group_by_class(cost_df, "account_name", "amount")
        if rows:
            _add_exec_table(doc,
                ["계정명", "건수", "금액", "IT", "보안"],
                rows, col_widths=[5, 2, 3, 2, 2], header_color=C_PURPLE
            )
        else:
            _add_body(doc, "비용 계정별 데이터가 없습니다.", color=C_GRAY)
    else:
        _add_body(doc, "비용 분석 데이터가 없습니다.", color=C_GRAY)

    doc.add_page_break()

    # ═══════════════════════════════════
    # ④ 주요 투자 항목 Top 10
    # ═══════════════════════════════════
    _add_section_heading(doc, "04", "주요 투자 항목 분석 (Top 10)")

    _add_sub_heading(doc, "자산 주요 항목 (감가상각비 기준)")
    top_a = _get_top_items(asset_df, "depreciation", 10)
    if top_a:
        _add_exec_table(doc,
            ["Description", "감가상각비", "분류", "매칭유형"],
            top_a, col_widths=[6.5, 3, 2.5, 3]
        )
    else:
        _add_body(doc, "자산 데이터가 없습니다.", color=C_GRAY)

    _add_sub_heading(doc, "비용 주요 항목 (금액 기준)")
    top_c = _get_top_items(cost_df, "amount", 10)
    if top_c:
        _add_exec_table(doc,
            ["Description", "금액", "분류", "매칭유형"],
            top_c, col_widths=[6.5, 3, 2.5, 3]
        )
    else:
        _add_body(doc, "비용 데이터가 없습니다.", color=C_GRAY)

    doc.add_page_break()

    # ═══════════════════════════════════
    # ⑤ 분류 기준 및 판단 로직
    # ═══════════════════════════════════
    _add_section_heading(doc, "05", "분류 기준 및 판단 로직 (Audit 대응)")

    _add_body(doc,
        "본 분류는 KISA 「정보보호 공시 가이드라인(2024.03)」의 자산분류표(p.72~77)를 기준으로 합니다."
    )

    _add_sub_heading(doc, "분류 원칙")
    _add_body(doc, "• 정보기술부문 — 포괄적 관점에서 작성 (전기통신설비, H/W·S/W 및 간접설비)", indent=True)
    _add_body(doc, "• 정보보호부문 — 보수적 관점에서 작성 (정보의 훼손·변조·유출 방지 수단)", indent=True)
    _add_body(doc, "• 제외 — 정보기술·정보보호와 직접적 관련이 없는 항목", indent=True)

    _add_sub_heading(doc, "자동 분류 로직 (4단계)")
    _add_exec_table(doc,
        ["단계", "방식", "설명", "점수 기준"],
        [
            ["1차", "완전일치", "기준정보 Description과 전처리 후 100% 동일", "100%"],
            ["2차", "포함매칭", "기준정보 Description 전체가 분석대상에 포함", "포함비율"],
            ["3차", "유사매칭", "RapidFuzz token_sort_ratio 유사도", "85% 이상"],
            ["4차", "AI 분류", "OpenAI + KISA 가이드라인 기준 자동 판단", "신뢰도"],
        ],
        col_widths=[2, 3, 6.5, 3],
        header_color=C_LIGHT_NAVY
    )

    _add_sub_heading(doc, "주요 판단 예시")
    examples = [
        ["CCTV", "보안 목적이면 정보보호, 일반 감시이면 제외"],
        ["백업 시스템", "일반 운영이면 IT, 원격지 재해복구이면 정보보호 포함"],
        ["VPN 장비", "IT + 정보보호 모두 해당 (이중 분류)"],
        ["방화벽/IPS", "정보보호 (IT에도 포함)"],
        ["사무용 프린터", "제외 (IT/보안 무관)"],
    ]
    _add_exec_table(doc,
        ["항목", "판단 기준"],
        examples, col_widths=[4, 12], header_color=C_GRAY
    )

    doc.add_page_break()

    # ═══════════════════════════════════
    # ⑥ 회계 정합성
    # ═══════════════════════════════════
    _add_section_heading(doc, "06", "회계 정합성 및 산출 방법")

    _add_exec_table(doc,
        ["항목", "내용"],
        [
            ["데이터 출처", "ERP 자산대장 및 비용원장"],
            ["자산 산출 기준", "연중 감가상각비 (발생주의)"],
            ["비용 산출 기준", "발생액 기준 (발생주의)"],
            ["중복 제거 로직", "IT ⊃ 정보보호 구조 적용 (정보보호는 IT에 포함)"],
            ["셰어드 서비스", "가이드라인 p.44~45 참조하여 배분 기준 적용"],
        ],
        col_widths=[4, 12],
        header_color=C_LIGHT_NAVY
    )

    _add_body(doc,
        "※ 2024 가이드라인은 회계 정합성 검토를 강조하며, "
        "자산대장·비용원장과 재무제표(감사보고서) 간의 정합성 확인이 필요합니다.",
        size=9, color=C_GRAY
    )

    doc.add_paragraph("")

    # ═══════════════════════════════════
    # ⑦ 매칭 결과 통계
    # ═══════════════════════════════════
    _add_section_heading(doc, "07", "분석 매칭 결과 통계")

    _add_sub_heading(doc, "자산 매칭 통계")
    a_stats = _match_type_stats(asset_df)
    if a_stats:
        _add_exec_table(doc, ["매칭유형", "건수", "비율"], a_stats, col_widths=[5, 5, 5])
    else:
        _add_body(doc, "자산 분석 데이터가 없습니다.", color=C_GRAY)

    _add_sub_heading(doc, "비용 매칭 통계")
    c_stats = _match_type_stats(cost_df)
    if c_stats:
        _add_exec_table(doc, ["매칭유형", "건수", "비율"], c_stats, col_widths=[5, 5, 5])
    else:
        _add_body(doc, "비용 분석 데이터가 없습니다.", color=C_GRAY)

    doc.add_page_break()

    # ═══════════════════════════════════
    # ⑧ 리스크 및 개선사항
    # ═══════════════════════════════════
    _add_section_heading(doc, "08", "리스크 및 개선사항")

    risks = _analyze_risks(stats, asset_df, cost_df, ratio)
    if risks:
        _add_exec_table(doc,
            ["항목", "현황", "개선방안"],
            risks, col_widths=[4, 6, 6]
        )
    else:
        _add_body(doc, "특이사항 없음")

    doc.add_paragraph("")

    # ── 참고 문헌 ──
    _add_section_heading(doc, "", "참고")
    _add_body(doc, "• KISA 정보보호 공시 가이드라인(2024.03 개정본)", indent=True)
    _add_body(doc, "  https://isds.kisa.or.kr/kr/bbs/view.do?bbsId=B0000011&menuNo=204948&nttId=38",
              size=9, color=C_ACCENT, indent=True)
    _add_body(doc, "• 「정보보호산업의 진흥에 관한 법률」 제13조", indent=True)
    _add_body(doc, "• 「정보보호 공시에 관한 고시」", indent=True)

    doc.add_paragraph("")
    _add_body(doc, f"본 보고서는 {today.strftime('%Y년 %m월 %d일')} AI기반 정보보호공시 솔루션에 의해 자동 생성되었습니다.",
              size=8, color=C_GRAY, italic=True)

    # ── 저장 ──
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ═══════════════════════════════════
# 데이터 집계 헬퍼 함수
# ═══════════════════════════════════

def _compute_stats(asset_df, cost_df):
    s = {}
    for prefix, df, amt_col in [("asset", asset_df, "depreciation"), ("cost", cost_df, "amount")]:
        if not df.empty and amt_col in df.columns:
            df[amt_col] = pd.to_numeric(df[amt_col], errors='coerce').fillna(0)
            s[f"{prefix}_it_total"] = abs(float(df.loc[df["it_yn"] == "O", amt_col].sum()))
            s[f"{prefix}_it_cnt"] = int((df["it_yn"] == "O").sum())
            s[f"{prefix}_sec_total"] = abs(float(df.loc[df["sec_yn"] == "O", amt_col].sum()))
            s[f"{prefix}_sec_cnt"] = int((df["sec_yn"] == "O").sum())
            s[f"{prefix}_exc_total"] = abs(float(df.loc[df["exclude_yn"] == "O", amt_col].sum()))
            s[f"{prefix}_exc_cnt"] = int((df["exclude_yn"] == "O").sum())
        else:
            for k in ["it_total","it_cnt","sec_total","sec_cnt","exc_total","exc_cnt"]:
                s[f"{prefix}_{k}"] = 0
    return s


def _group_by_class(df, class_col, amount_col):
    if class_col not in df.columns or amount_col not in df.columns:
        return []
    df = df.copy()
    df[amount_col] = pd.to_numeric(df[amount_col], errors='coerce').fillna(0)
    grouped = df.groupby(class_col).agg(
        count=(amount_col, 'size'),
        total=(amount_col, lambda x: abs(x.sum())),
        it_cnt=('it_yn', lambda x: (x == 'O').sum()),
        sec_cnt=('sec_yn', lambda x: (x == 'O').sum()),
    ).reset_index().sort_values('total', ascending=False)

    rows = []
    for _, r in grouped.head(15).iterrows():
        rows.append([str(r[class_col]), f"{int(r['count'])}건",
                     _fmt_won(r['total']), f"{int(r['it_cnt'])}건", f"{int(r['sec_cnt'])}건"])
    return rows


def _get_top_items(df, amount_col, n=10):
    if df.empty or amount_col not in df.columns:
        return []
    df = df.copy()
    df[amount_col] = pd.to_numeric(df[amount_col], errors='coerce').fillna(0)
    df["_abs"] = df[amount_col].abs()
    top = df.nlargest(n, "_abs")
    rows = []
    for _, r in top.iterrows():
        desc = str(r.get("description", ""))[:40]
        cls = []
        if r.get("it_yn") == "O": cls.append("IT")
        if r.get("sec_yn") == "O": cls.append("보안")
        if r.get("exclude_yn") == "O": cls.append("제외")
        rows.append([desc, _fmt_won(r[amount_col]), "/".join(cls) or "-", str(r.get("match_type", ""))])
    return rows


def _match_type_stats(df):
    if df.empty or "match_type" not in df.columns:
        return []
    counts = df["match_type"].value_counts()
    total = len(df)
    return [[str(mt), f"{cnt}건", f"{cnt/total*100:.1f}%"] for mt, cnt in counts.items() if mt]


def _analyze_risks(stats, asset_df, cost_df, ratio):
    risks = []
    if ratio < 5:
        risks.append(["보안 투자 비율 미흡", f"IT 대비 {ratio:.1f}% (업계 권장 5~15%)", "정보보호 예산 확대 검토 필요"])
    elif ratio < 10:
        risks.append(["보안 투자 비율 보통", f"IT 대비 {ratio:.1f}%", "업계 평균 수준, 강화 검토"])
    else:
        risks.append(["보안 투자 비율 양호", f"IT 대비 {ratio:.1f}%", "현 수준 유지"])

    for label, df in [("자산", asset_df), ("비용", cost_df)]:
        if not df.empty and "match_type" in df.columns:
            unm = (df["match_type"] == "미매칭").sum()
            if unm > 0:
                risks.append([f"{label} 미분류 항목", f"{unm}건 미매칭", "기준정보 보완 또는 수동 검토"])
            ai_c = (df["match_type"] == "AI분류").sum()
            if ai_c > 0:
                risks.append([f"{label} AI 자동분류", f"{ai_c}건 (수동 검증 권장)", "AI 판단 결과 확인 필요"])

    return risks if risks else [["특이사항 없음", "-", "-"]]


# ═══════════════════════════════════
# 미리보기용 데이터 생성
# ═══════════════════════════════════

def compute_preview_data(asset_df: pd.DataFrame, cost_df: pd.DataFrame) -> dict:
    """리포트 미리보기에 필요한 모든 데이터를 계산하여 반환"""
    stats = _compute_stats(asset_df, cost_df)
    it_total = stats["asset_it_total"] + stats["cost_it_total"]
    sec_total = stats["asset_sec_total"] + stats["cost_sec_total"]
    ratio = (sec_total / it_total * 100) if it_total > 0 else 0

    # 자산 분류별
    asset_by_class = []
    if not asset_df.empty and "asset_class" in asset_df.columns:
        asset_by_class = _group_by_class(asset_df, "asset_class", "depreciation")

    # 비용 계정별
    cost_by_acct = []
    if not cost_df.empty and "account_name" in cost_df.columns:
        cost_by_acct = _group_by_class(cost_df, "account_name", "amount")

    # Top 10
    top_assets = _get_top_items(asset_df, "depreciation", 10)
    top_costs = _get_top_items(cost_df, "amount", 10)

    # 매칭 통계
    asset_match = _match_type_stats(asset_df)
    cost_match = _match_type_stats(cost_df)

    # 리스크
    risks = _analyze_risks(stats, asset_df, cost_df, ratio)

    return {
        "stats": stats,
        "it_total": it_total,
        "sec_total": sec_total,
        "ratio": ratio,
        "asset_by_class": asset_by_class,
        "cost_by_acct": cost_by_acct,
        "top_assets": top_assets,
        "top_costs": top_costs,
        "asset_match": asset_match,
        "cost_match": cost_match,
        "risks": risks,
    }
